"""
将 HTML 转换为 Word (.docx) 文件。
使用 python-docx 生成文档，通过 lxml 解析 HTML 结构。
"""
import io
import re
import base64
import logging
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import html as lxml_html

logger = logging.getLogger(__name__)

_UPLOAD_DIR = Path(__file__).resolve().parent.parent.parent / "uploads" / "images"


def _resolve_image_src(src: str) -> bytes | None:
    """将 img src（data URI 或 /api/images/xxx）解析为图片字节。"""
    if not src:
        return None
    if src.startswith("data:"):
        m = re.match(r"data:image/[^;]+;base64,(.+)", src, re.DOTALL)
        if m:
            try:
                return base64.b64decode(m.group(1))
            except Exception:
                return None
    elif src.startswith("/api/images/"):
        filename = src.split("/")[-1]
        path = _UPLOAD_DIR / filename
        if path.exists():
            return path.read_bytes()
    return None


def _add_inline_runs(paragraph, element):
    """递归处理内联元素（加粗/斜体/链接/文本等）。"""
    tag = element.tag if hasattr(element, 'tag') else None

    if tag is None or isinstance(tag, str) is False:
        if element.text:
            paragraph.add_run(element.text)
        return

    tag_lower = tag.lower() if isinstance(tag, str) else ""

    if element.text:
        run = paragraph.add_run(element.text)
        if tag_lower in ('strong', 'b'):
            run.bold = True
        elif tag_lower in ('em', 'i'):
            run.italic = True
        elif tag_lower == 'u':
            run.underline = True

    for child in element:
        child_tag = (child.tag or "").lower() if hasattr(child, 'tag') and isinstance(child.tag, str) else ""
        if child_tag in ('strong', 'b'):
            r = paragraph.add_run(child.text_content())
            r.bold = True
        elif child_tag in ('em', 'i'):
            r = paragraph.add_run(child.text_content())
            r.italic = True
        elif child_tag == 'u':
            r = paragraph.add_run(child.text_content())
            r.underline = True
        elif child_tag == 'br':
            paragraph.add_run('\n')
        elif child_tag == 'img':
            _add_image_to_paragraph(paragraph, child)
        else:
            _add_inline_runs(paragraph, child)

        if child.tail:
            run = paragraph.add_run(child.tail)
            if tag_lower in ('strong', 'b'):
                run.bold = True
            elif tag_lower in ('em', 'i'):
                run.italic = True


def _add_image_to_paragraph(paragraph, img_el):
    """将 <img> 元素插入到段落中。"""
    src = img_el.get("src", "")
    img_bytes = _resolve_image_src(src)
    if not img_bytes:
        return
    try:
        stream = io.BytesIO(img_bytes)
        run = paragraph.add_run()
        run.add_picture(stream, width=Inches(5.5))
    except Exception:
        logger.warning("docx_export: failed to add image, src=%s", src[:80])


def _get_text(el) -> str:
    """获取元素的纯文本内容。"""
    return (el.text_content() or "").strip()


def _process_element(doc: Document, el, list_level: int = 0):
    """递归处理 HTML 元素，写入 docx Document。"""
    tag = (el.tag or "").lower() if hasattr(el, 'tag') and isinstance(el.tag, str) else ""

    if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
        level = int(tag[1])
        p = doc.add_heading(level=level)
        _add_inline_runs(p, el)

    elif tag == 'p':
        has_img = el.find('.//img') is not None
        p = doc.add_paragraph()
        if has_img:
            for child in el:
                child_tag = (child.tag or "").lower() if hasattr(child, 'tag') and isinstance(child.tag, str) else ""
                if child_tag == 'img':
                    _add_image_to_paragraph(p, child)
                else:
                    _add_inline_runs(p, child)
                if child.tail:
                    p.add_run(child.tail)
            if el.text:
                run = p.add_run(el.text)
                run_copy = p.runs[0] if p.runs else None
                if run_copy:
                    p.runs.insert(0, p.runs.pop())
        else:
            _add_inline_runs(p, el)

    elif tag in ('ul', 'ol'):
        for li in el:
            li_tag = (li.tag or "").lower() if hasattr(li, 'tag') and isinstance(li.tag, str) else ""
            if li_tag == 'li':
                p = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
                _add_inline_runs(p, li)

    elif tag == 'table':
        _process_table(doc, el)

    elif tag == 'blockquote':
        text = _get_text(el)
        if text:
            p = doc.add_paragraph(text, style='Quote')

    elif tag == 'pre':
        text = _get_text(el)
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    elif tag == 'img':
        p = doc.add_paragraph()
        _add_image_to_paragraph(p, el)

    elif tag in ('div', 'section', 'article', 'main', 'body'):
        for child in el:
            _process_element(doc, child, list_level)

    elif tag == 'br':
        doc.add_paragraph()

    else:
        text = _get_text(el)
        if text:
            doc.add_paragraph(text)


def _process_table(doc: Document, table_el):
    """将 HTML <table> 转为 docx 表格。"""
    rows_data: list[list[str]] = []
    for tr in table_el.iter('tr'):
        cells = []
        for td in tr:
            td_tag = (td.tag or "").lower() if hasattr(td, 'tag') and isinstance(td.tag, str) else ""
            if td_tag in ('td', 'th'):
                cells.append(_get_text(td))
        if cells:
            rows_data.append(cells)

    if not rows_data:
        return

    max_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=max_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(rows_data):
        for j, cell_text in enumerate(row):
            if j < max_cols:
                table.rows[i].cells[j].text = cell_text


def html_to_docx(html_content: str) -> io.BytesIO:
    """
    将 HTML 字符串转换为 .docx 文件，返回 BytesIO。
    """
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)

    wrapped = f"<body>{html_content}</body>"
    try:
        tree = lxml_html.fromstring(wrapped)
    except Exception:
        logger.exception("docx_export: failed to parse HTML")
        p = doc.add_paragraph(html_content)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    for child in tree:
        _process_element(doc, child)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
